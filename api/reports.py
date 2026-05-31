from http.server import BaseHTTPRequestHandler
import io
import re
import zipfile
import datetime
import sys
import os
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from api._lib.db import query
from api._lib.auth import get_user_from_request
from api._lib.response import error_response, options_response, get_query_params

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

PASS_THRESHOLD = 70.0


def _safe(name):
    return re.sub(r'[^\w\-]+', '_', str(name))


def _build_report(test_id):
    """读取试验数据，生成 .docx 字节流（论文 5.8 报告结构）。"""
    test = query(
        """SELECT t.*, p.name as project_name, p.pipe_diameter, p.layer_width,
                  p.layer_thickness, p.location
           FROM tests t LEFT JOIN projects p ON t.project_id = p.id
           WHERE t.id = %s""", (test_id,), fetchone=True)
    if not test:
        return None, None

    strips = query(
        """SELECT strip_number,
                  ROUND(AVG(force_value)::numeric,2) as avg_force,
                  ROUND(MAX(force_value)::numeric,2) as max_force,
                  ROUND(MIN(force_value)::numeric,2) as min_force,
                  ROUND(STDDEV(force_value)::numeric,2) as std_force,
                  (AVG(force_value) >= %s) as pass_fail
           FROM data_points WHERE test_id = %s
           GROUP BY strip_number ORDER BY strip_number""",
        (PASS_THRESHOLD, test_id), fetchall=True)

    overall = query(
        """SELECT ROUND(AVG(force_value)::numeric,2) as avg_force,
                  ROUND(MAX(force_value)::numeric,2) as max_force,
                  ROUND(MIN(force_value)::numeric,2) as min_force,
                  COUNT(*) as total_points,
                  ROUND(AVG(CASE WHEN force_value>=%s THEN 1.0 ELSE 0.0 END)*100,2) as pass_rate
           FROM data_points WHERE test_id = %s""",
        (PASS_THRESHOLD, test_id), fetchone=True)

    doc = Document()
    title = doc.add_heading('管道补口防腐层剥离试验报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph('Pipeline Joint-Coating Peeling Test Report')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('一、项目与试样信息', level=1)
    info = doc.add_table(rows=0, cols=4)
    info.style = 'Light Grid Accent 1'
    rows = [
        ('项目名称', test.get('project_name') or '-', '试验编号', test.get('test_number') or '-'),
        ('试样编号', test.get('sample_name') or '-', '操作员', test.get('operator') or '-'),
        ('管道直径(mm)', str(test.get('pipe_diameter') or '-'), '防腐层宽度(mm)', str(test.get('layer_width') or '-')),
        ('剥离速度(mm/min)', str(test.get('peel_speed') or '-'), '条带数量', str(test.get('n_strips') or len(strips))),
        ('检测地点', test.get('location') or '-', '试验状态', test.get('status') or '-'),
    ]
    for a, b, c, d in rows:
        cells = info.add_row().cells
        cells[0].text, cells[1].text, cells[2].text, cells[3].text = a, b, c, d

    doc.add_heading('二、关键指标摘要', level=1)
    p = doc.add_paragraph()
    p.add_run(f"整体平均剥离力：{overall['avg_force']} N\n")
    p.add_run(f"峰值剥离力：{overall['max_force']} N\n")
    p.add_run(f"最小剥离力：{overall['min_force']} N\n")
    p.add_run(f"力-位采样点总数：{overall['total_points']}\n")
    p.add_run(f"合格阈值：{PASS_THRESHOLD:.0f} N（≈35 N/cm）\n")
    p.add_run(f"采样点合格率：{overall['pass_rate']} %")

    doc.add_heading('三、各条带剥离力统计', level=1)
    tbl = doc.add_table(rows=1, cols=6)
    tbl.style = 'Light Grid Accent 1'
    hdr = tbl.rows[0].cells
    for i, h in enumerate(['条带', '平均(N)', '峰值(N)', '最小(N)', '标准差', '判定']):
        hdr[i].text = h
    pass_n = 0
    for s in strips:
        c = tbl.add_row().cells
        c[0].text = str(s['strip_number'])
        c[1].text = str(s['avg_force'])
        c[2].text = str(s['max_force'])
        c[3].text = str(s['min_force'])
        c[4].text = str(s['std_force'])
        ok = bool(s['pass_fail'])
        pass_n += 1 if ok else 0
        c[5].text = '合格' if ok else '不合格'

    doc.add_heading('四、合格判定结论', level=1)
    rate = (100.0 * pass_n / len(strips)) if strips else 0.0
    concl = doc.add_paragraph()
    run = concl.add_run(
        f"本次试验共测量 {len(strips)} 条剥离条带，其中合格 {pass_n} 条，"
        f"条带合格率 {rate:.1f}%。" +
        ("整体粘接质量良好。" if rate >= 60 else "存在较多弱粘/缺陷区域，建议复检该补口。"))
    run.bold = True

    foot = doc.add_paragraph(
        f"报告生成时间：{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}    "
        "管道补口自动剥离系统 PCS")
    foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    fname = f"report_{_safe(test.get('test_number') or test_id)}.docx"
    return buf.getvalue(), fname


class handler(BaseHTTPRequestHandler):
    def do_OPTIONS(self):
        options_response(self)

    def do_GET(self):
        try:
            params = get_query_params(self)
            action = params.get('action', 'test')
            payload = get_user_from_request(self.headers)

            if action == 'test':
                test_id = params.get('test_id') or params.get('id')
                if not test_id:
                    error_response(self, '缺少试验ID')
                    return
                content, fname = _build_report(test_id)
                if content is None:
                    error_response(self, '试验不存在', 404)
                    return
                if payload:
                    query("INSERT INTO audit_log (user_id, action, resource_type, resource_id) VALUES (%s,%s,%s,%s)",
                          (payload['user_id'], 'export_report', 'test', int(test_id)))
                self._send_file(content, fname,
                                'application/vnd.openxmlformats-officedocument.wordprocessingml.document')

            elif action == 'project':
                project_id = params.get('project_id') or params.get('id')
                if not project_id:
                    error_response(self, '缺少项目ID')
                    return
                tests = query(
                    "SELECT id FROM tests WHERE project_id = %s AND status = 'completed' ORDER BY id",
                    (project_id,), fetchall=True)
                if not tests:
                    error_response(self, '该项目暂无已完成试验', 404)
                    return
                zbuf = io.BytesIO()
                with zipfile.ZipFile(zbuf, 'w', zipfile.ZIP_DEFLATED) as zf:
                    for t in tests:
                        content, fname = _build_report(t['id'])
                        if content:
                            zf.writestr(fname, content)
                zbuf.seek(0)
                if payload:
                    query("INSERT INTO audit_log (user_id, action, resource_type, resource_id) VALUES (%s,%s,%s,%s)",
                          (payload['user_id'], 'export_reports_zip', 'project', int(project_id)))
                self._send_file(zbuf.getvalue(), f'reports_project_{project_id}.zip',
                                'application/zip')
            else:
                error_response(self, f'未知操作: {action}')
        except Exception as e:
            error_response(self, str(e), 500)

    def _send_file(self, content, filename, content_type):
        self.send_response(200)
        self.send_header('Content-Type', content_type)
        self.send_header('Content-Disposition', f'attachment; filename={filename}')
        self.send_header('Content-Length', str(len(content)))
        self.send_header('Access-Control-Allow-Origin', '*')
        self.end_headers()
        self.wfile.write(content)
